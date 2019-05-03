import os
import comtypes.client
from docx import Document


class Docx:
    def __init__(self, docx_path: str):
        if os.path.exists(docx_path):
            self._in_path = os.path.abspath(docx_path)
            self.doc_obj = Document(self._in_path)
        else:
            raise FileExistsError

    def get_path(self):
        return self._in_path

    def update_path(self, new_path: str):
        self._in_path = os.path.abspath(new_path)
        return self.get_path()

    def replace(self, template_str, value_str):
        self.static_replace_substring(self.doc_obj, template_str, value_str)

    def save(self):
        self.doc_obj.save(self._in_path)

    def to_pdf(self):
        tempfile_path = os.path.abspath('temp_pdf.temp')
        folder, file = os.path.split(self._in_path)
        f_name, f_ext = os.path.splitext(file)
        pdf_path = os.path.join(folder, (f_name + '.pdf'))
        self.doc_obj.save(tempfile_path)
        self.static_to_pdf(tempfile_path, pdf_path)
        os.remove(tempfile_path)

    @staticmethod
    def static_to_pdf(in_file: str, out_file: str):

        print('Creating PDF...')
        in_path = os.path.abspath(in_file)
        out_path = os.path.abspath(out_file)
        wd_format_pdf = 17

        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(in_path)
        doc.SaveAs(out_path, FileFormat=wd_format_pdf)
        doc.Close()
        word.Quit()
        print('Done')

    @staticmethod
    def static_replace_substring(doc_obj, template_str, value_str, max_cnt=100):
        was_found = True
        while (max_cnt > 0) and was_found:
            was_found = False
            for p in doc_obj.paragraphs:
                pos = p.text.find(template_str)
                if pos >= 0:
                    was_found = True
                    inline = p.runs
                    pos_end = len(template_str)
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if pos_end == 0:
                            # replace completed
                            break
                        if (pos_end > 0) and (pos_end < len(template_str)):
                            # part of srting in another run, delete from start
                            # length of part that in this run
                            part_len = min(len(inline[i].text), pos_end)
                            inline[i].text = inline[i].text[part_len:]
                            pos_end = pos_end - part_len
                        if pos - len(inline[i].text) < 0:
                            # Use slicing to extract those parts of the original string to be kept
                            # length of part that in this run
                            part_len = min(
                                len(inline[i].text) - pos, len(template_str))
                            inline[i].text = inline[i].text[:pos] + \
                                             value_str + inline[i].text[(pos + part_len):]
                            pos_end = pos_end - part_len
                        else:
                            pos = pos - len(inline[i].text)
            max_cnt -= 1

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    Docx.static_replace_substring(cell, template_str, value_str, max_cnt)
